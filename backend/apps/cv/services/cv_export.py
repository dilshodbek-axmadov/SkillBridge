"""
CV Export Service
=================
Generates PDF and DOCX files from CV data.
Uses ReportLab for PDF and python-docx for DOCX.
"""

import io
from apps.cv.models import CV


class CVExportService:
    """Export CV to PDF or DOCX format."""

    def __init__(self, cv: CV):
        self.cv = cv
        self.sections = list(
            cv.cv_sections.filter(is_visible=True).order_by('display_order')
        )

    def export_pdf(self):
        """
        Generate PDF from CV data using ReportLab.
        Returns bytes buffer.
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table,
            TableStyle, HRFlowable,
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20 * mm,
            leftMargin=20 * mm,
            topMargin=15 * mm,
            bottomMargin=15 * mm,
        )

        styles = getSampleStyleSheet()
        # Custom styles
        styles.add(ParagraphStyle(
            'CVName',
            parent=styles['Title'],
            fontSize=22,
            textColor=HexColor('#1a1a2e'),
            spaceAfter=4,
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            'CVContact',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor('#555555'),
            alignment=TA_CENTER,
            spaceAfter=6,
        ))
        styles.add(ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=HexColor('#1a1a2e'),
            spaceBefore=10,
            spaceAfter=4,
            borderWidth=0,
        ))
        styles.add(ParagraphStyle(
            'CVBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=4,
        ))
        styles.add(ParagraphStyle(
            'CVBold',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=2,
        ))
        styles.add(ParagraphStyle(
            'CVSmall',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor('#666666'),
            spaceAfter=2,
        ))

        elements = []
        primary_color = HexColor('#1a1a2e')

        for section in self.sections:
            content = section.content or {}
            section_type = section.section_type

            if section_type == 'personal_info':
                elements.extend(
                    self._pdf_personal_info(content, styles)
                )
                elements.append(HRFlowable(
                    width="100%", thickness=1,
                    color=primary_color, spaceAfter=8,
                ))

            elif section_type == 'summary':
                elements.append(Paragraph('PROFESSIONAL SUMMARY', styles['SectionTitle']))
                elements.append(HRFlowable(
                    width="100%", thickness=0.5,
                    color=HexColor('#cccccc'), spaceAfter=6,
                ))
                text = content.get('text', '')
                if text:
                    elements.append(Paragraph(text, styles['CVBody']))

            elif section_type == 'experience':
                elements.extend(
                    self._pdf_experience(content, styles)
                )

            elif section_type == 'education':
                elements.extend(
                    self._pdf_education(content, styles)
                )

            elif section_type == 'skills':
                elements.extend(
                    self._pdf_skills(content, styles)
                )

            elif section_type == 'projects':
                elements.extend(
                    self._pdf_projects(content, styles)
                )

            elif section_type == 'certifications':
                elements.extend(
                    self._pdf_certifications(content, styles)
                )

            elif section_type == 'languages':
                elements.extend(
                    self._pdf_languages(content, styles)
                )

            elif section_type == 'awards':
                elements.extend(
                    self._pdf_awards(content, styles)
                )

            elements.append(Spacer(1, 4))

        if elements:
            doc.build(elements)

        buffer.seek(0)
        return buffer

    def _pdf_personal_info(self, content, styles):
        """Render personal info section for PDF."""
        from reportlab.platypus import Paragraph, Spacer

        elements = []
        name = content.get('full_name', '')
        if name:
            elements.append(Paragraph(name, styles['CVName']))

        contact_parts = []
        if content.get('email'):
            contact_parts.append(content['email'])
        if content.get('phone'):
            contact_parts.append(content['phone'])
        if content.get('location'):
            contact_parts.append(content['location'])
        if contact_parts:
            elements.append(Paragraph(
                ' | '.join(contact_parts), styles['CVContact']
            ))

        link_parts = []
        if content.get('github_url'):
            link_parts.append(content['github_url'])
        if content.get('linkedin_url'):
            link_parts.append(content['linkedin_url'])
        if content.get('portfolio_url'):
            link_parts.append(content['portfolio_url'])
        if link_parts:
            elements.append(Paragraph(
                ' | '.join(link_parts), styles['CVContact']
            ))

        elements.append(Spacer(1, 4))
        return elements

    def _pdf_experience(self, content, styles):
        """Render experience section for PDF."""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        positions = content.get('positions', [])
        if not positions:
            return elements

        elements.append(Paragraph('WORK EXPERIENCE', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for pos in positions:
            title = pos.get('title', '')
            company = pos.get('company', '')
            location = pos.get('location', '')
            start = pos.get('start_date', '')
            end = 'Present' if pos.get('current') else pos.get('end_date', '')
            date_range = f"{start} - {end}" if start else ''

            elements.append(Paragraph(
                f"<b>{title}</b> — {company}", styles['CVBold']
            ))
            if location or date_range:
                elements.append(Paragraph(
                    f"{location}  {date_range}".strip(), styles['CVSmall']
                ))

            for resp in pos.get('responsibilities', []):
                elements.append(Paragraph(f"• {resp}", styles['CVBody']))
            for ach in pos.get('achievements', []):
                elements.append(Paragraph(f"★ {ach}", styles['CVBody']))

            elements.append(Spacer(1, 4))

        return elements

    def _pdf_education(self, content, styles):
        """Render education section for PDF."""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        degrees = content.get('degrees', [])
        if not degrees:
            return elements

        elements.append(Paragraph('EDUCATION', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for deg in degrees:
            degree = deg.get('degree', '')
            field = deg.get('field', '')
            institution = deg.get('institution', '')
            start = deg.get('start_date', '')
            end = deg.get('end_date', '')
            date_range = f"{start} - {end}" if start else ''

            elements.append(Paragraph(
                f"<b>{degree} in {field}</b>", styles['CVBold']
            ))
            elements.append(Paragraph(
                f"{institution}  {date_range}".strip(), styles['CVSmall']
            ))
            if deg.get('gpa'):
                elements.append(Paragraph(
                    f"GPA: {deg['gpa']}", styles['CVSmall']
                ))
            elements.append(Spacer(1, 4))

        return elements

    def _pdf_skills(self, content, styles):
        """Render skills section for PDF."""
        from reportlab.platypus import Paragraph
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        categories = content.get('categories', [])
        if not categories:
            return elements

        elements.append(Paragraph('SKILLS', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for cat in categories:
            name = cat.get('name', '')
            skills = ', '.join(cat.get('skills', []))
            elements.append(Paragraph(
                f"<b>{name}:</b> {skills}", styles['CVBody']
            ))

        return elements

    def _pdf_projects(self, content, styles):
        """Render projects section for PDF."""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        projects = content.get('projects', [])
        if not projects:
            return elements

        elements.append(Paragraph('PROJECTS', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for proj in projects:
            name = proj.get('name', '')
            desc = proj.get('description', '')
            techs = ', '.join(proj.get('technologies', []))

            elements.append(Paragraph(f"<b>{name}</b>", styles['CVBold']))
            if desc:
                elements.append(Paragraph(desc, styles['CVBody']))
            if techs:
                elements.append(Paragraph(
                    f"<i>Technologies: {techs}</i>", styles['CVSmall']
                ))

            links = []
            if proj.get('github_url'):
                links.append(proj['github_url'])
            if proj.get('live_url'):
                links.append(proj['live_url'])
            if links:
                elements.append(Paragraph(
                    ' | '.join(links), styles['CVSmall']
                ))

            for hl in proj.get('highlights', []):
                elements.append(Paragraph(f"• {hl}", styles['CVBody']))

            elements.append(Spacer(1, 4))

        return elements

    def _pdf_certifications(self, content, styles):
        """Render certifications section for PDF."""
        from reportlab.platypus import Paragraph
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        certs = content.get('certifications', [])
        if not certs:
            return elements

        elements.append(Paragraph('CERTIFICATIONS', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for cert in certs:
            name = cert.get('name', '')
            issuer = cert.get('issuer', '')
            date = cert.get('date', '')
            elements.append(Paragraph(
                f"<b>{name}</b> — {issuer} ({date})", styles['CVBody']
            ))

        return elements

    def _pdf_languages(self, content, styles):
        """Render languages section for PDF."""
        from reportlab.platypus import Paragraph
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        languages = content.get('languages', [])
        if not languages:
            return elements

        elements.append(Paragraph('LANGUAGES', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for lang in languages:
            name = lang.get('language', '')
            prof = lang.get('proficiency', '')
            elements.append(Paragraph(
                f"<b>{name}</b> — {prof}", styles['CVBody']
            ))

        return elements

    def _pdf_awards(self, content, styles):
        """Render awards section for PDF."""
        from reportlab.platypus import Paragraph
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import HRFlowable

        elements = []
        awards = content.get('awards', [])
        if not awards:
            return elements

        elements.append(Paragraph('AWARDS & ACHIEVEMENTS', styles['SectionTitle']))
        elements.append(HRFlowable(
            width="100%", thickness=0.5,
            color=HexColor('#cccccc'), spaceAfter=6,
        ))

        for award in awards:
            title = award.get('title', '')
            issuer = award.get('issuer', '')
            date = award.get('date', '')
            desc = award.get('description', '')
            elements.append(Paragraph(
                f"<b>{title}</b> — {issuer} ({date})", styles['CVBody']
            ))
            if desc:
                elements.append(Paragraph(desc, styles['CVSmall']))

        return elements

    def export_docx(self):
        """
        Generate DOCX from CV data using python-docx.
        Returns bytes buffer.
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section_doc in doc.sections:
            section_doc.top_margin = Inches(0.6)
            section_doc.bottom_margin = Inches(0.6)
            section_doc.left_margin = Inches(0.8)
            section_doc.right_margin = Inches(0.8)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        primary_color = RGBColor(0x1a, 0x1a, 0x2e)

        for section in self.sections:
            content = section.content or {}
            section_type = section.section_type

            if section_type == 'personal_info':
                self._docx_personal_info(doc, content, primary_color)
            elif section_type == 'summary':
                self._docx_section_heading(doc, 'PROFESSIONAL SUMMARY', primary_color)
                text = content.get('text', '')
                if text:
                    doc.add_paragraph(text)
            elif section_type == 'experience':
                self._docx_experience(doc, content, primary_color)
            elif section_type == 'education':
                self._docx_education(doc, content, primary_color)
            elif section_type == 'skills':
                self._docx_skills(doc, content, primary_color)
            elif section_type == 'projects':
                self._docx_projects(doc, content, primary_color)
            elif section_type == 'certifications':
                self._docx_certifications(doc, content, primary_color)
            elif section_type == 'languages':
                self._docx_languages(doc, content, primary_color)
            elif section_type == 'awards':
                self._docx_awards(doc, content, primary_color)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _docx_section_heading(self, doc, title, color):
        """Add a section heading with underline."""
        from docx.shared import Pt

        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = color
        p.space_after = Pt(4)

    def _docx_personal_info(self, doc, content, color):
        """Render personal info for DOCX."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        name = content.get('full_name', '')
        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = color

        contact_parts = []
        if content.get('email'):
            contact_parts.append(content['email'])
        if content.get('phone'):
            contact_parts.append(content['phone'])
        if content.get('location'):
            contact_parts.append(content['location'])
        if contact_parts:
            p = doc.add_paragraph(' | '.join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)

        link_parts = []
        if content.get('github_url'):
            link_parts.append(content['github_url'])
        if content.get('linkedin_url'):
            link_parts.append(content['linkedin_url'])
        if content.get('portfolio_url'):
            link_parts.append(content['portfolio_url'])
        if link_parts:
            p = doc.add_paragraph(' | '.join(link_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)

    def _docx_experience(self, doc, content, color):
        """Render experience section for DOCX."""
        from docx.shared import Pt

        positions = content.get('positions', [])
        if not positions:
            return

        self._docx_section_heading(doc, 'WORK EXPERIENCE', color)

        for pos in positions:
            title = pos.get('title', '')
            company = pos.get('company', '')
            location = pos.get('location', '')
            start = pos.get('start_date', '')
            end = 'Present' if pos.get('current') else pos.get('end_date', '')

            p = doc.add_paragraph()
            run = p.add_run(f"{title} — {company}")
            run.bold = True
            run.font.size = Pt(10)

            if location or start:
                date_range = f"{start} - {end}" if start else ''
                p = doc.add_paragraph(f"{location}  {date_range}".strip())
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = color

            for resp in pos.get('responsibilities', []):
                doc.add_paragraph(resp, style='List Bullet')
            for ach in pos.get('achievements', []):
                p = doc.add_paragraph(ach, style='List Bullet')

    def _docx_education(self, doc, content, color):
        """Render education section for DOCX."""
        from docx.shared import Pt

        degrees = content.get('degrees', [])
        if not degrees:
            return

        self._docx_section_heading(doc, 'EDUCATION', color)

        for deg in degrees:
            degree = deg.get('degree', '')
            field = deg.get('field', '')
            institution = deg.get('institution', '')
            start = deg.get('start_date', '')
            end = deg.get('end_date', '')

            p = doc.add_paragraph()
            run = p.add_run(f"{degree} in {field}")
            run.bold = True
            run.font.size = Pt(10)

            date_range = f"{start} - {end}" if start else ''
            p = doc.add_paragraph(f"{institution}  {date_range}".strip())
            for run in p.runs:
                run.font.size = Pt(9)

            if deg.get('gpa'):
                doc.add_paragraph(f"GPA: {deg['gpa']}")

    def _docx_skills(self, doc, content, color):
        """Render skills section for DOCX."""
        categories = content.get('categories', [])
        if not categories:
            return

        self._docx_section_heading(doc, 'SKILLS', color)

        for cat in categories:
            name = cat.get('name', '')
            skills = ', '.join(cat.get('skills', []))
            p = doc.add_paragraph()
            run = p.add_run(f"{name}: ")
            run.bold = True
            p.add_run(skills)

    def _docx_projects(self, doc, content, color):
        """Render projects section for DOCX."""
        from docx.shared import Pt

        projects = content.get('projects', [])
        if not projects:
            return

        self._docx_section_heading(doc, 'PROJECTS', color)

        for proj in projects:
            name = proj.get('name', '')
            desc = proj.get('description', '')
            techs = ', '.join(proj.get('technologies', []))

            p = doc.add_paragraph()
            run = p.add_run(name)
            run.bold = True

            if desc:
                doc.add_paragraph(desc)

            if techs:
                p = doc.add_paragraph()
                run = p.add_run(f"Technologies: {techs}")
                run.italic = True
                run.font.size = Pt(9)

            for hl in proj.get('highlights', []):
                doc.add_paragraph(hl, style='List Bullet')

    def _docx_certifications(self, doc, content, color):
        """Render certifications section for DOCX."""
        certs = content.get('certifications', [])
        if not certs:
            return

        self._docx_section_heading(doc, 'CERTIFICATIONS', color)

        for cert in certs:
            name = cert.get('name', '')
            issuer = cert.get('issuer', '')
            date = cert.get('date', '')
            p = doc.add_paragraph()
            run = p.add_run(f"{name}")
            run.bold = True
            p.add_run(f" — {issuer} ({date})")

    def _docx_languages(self, doc, content, color):
        """Render languages section for DOCX."""
        languages = content.get('languages', [])
        if not languages:
            return

        self._docx_section_heading(doc, 'LANGUAGES', color)

        for lang in languages:
            name = lang.get('language', '')
            prof = lang.get('proficiency', '')
            p = doc.add_paragraph()
            run = p.add_run(f"{name}")
            run.bold = True
            p.add_run(f" — {prof}")

    def _docx_awards(self, doc, content, color):
        """Render awards section for DOCX."""
        awards = content.get('awards', [])
        if not awards:
            return

        self._docx_section_heading(doc, 'AWARDS & ACHIEVEMENTS', color)

        for award in awards:
            title = award.get('title', '')
            issuer = award.get('issuer', '')
            date = award.get('date', '')
            desc = award.get('description', '')
            p = doc.add_paragraph()
            run = p.add_run(f"{title}")
            run.bold = True
            p.add_run(f" — {issuer} ({date})")
            if desc:
                doc.add_paragraph(desc)
