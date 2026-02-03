"""
DOCX Parser
===========
backend/apps/users/cv_parser/docx_parser.py

Extract text from DOCX files using python-docx.
"""

import logging
from pathlib import Path
from typing import Optional

from .base_parser import BaseParser

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """
    Parse DOCX files and extract text.
    
    Extracts from:
    - Paragraphs
    - Tables
    - Headers/Footers
    """
    
    def parse(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    
                    if row_data:
                        text_parts.append(' | '.join(row_data))
            
            # Extract from headers (if any)
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)
            
            combined_text = '\n'.join(text_parts)
            return self.clean_text(combined_text) if combined_text else None
        
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return None
    
    def validate(self, file_path: str) -> bool:
        """Validate DOCX file."""
        try:
            path = Path(file_path)
            
            # Check file exists
            if not path.exists():
                return False
            
            # Check extension
            if path.suffix.lower() not in ['.docx', '.doc']:
                return False
            
            # Check file size (max 5MB)
            if path.stat().st_size > 5 * 1024 * 1024:
                logger.warning("DOCX file too large (>5MB)")
                return False
            
            # Try to open with python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                
                # Check if document has content
                if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                    logger.warning("DOCX has no content")
                    return False
                
                return True
            
            except Exception as e:
                logger.error(f"DOCX validation failed: {e}")
                return False
        
        except Exception as e:
            logger.error(f"Error validating DOCX: {e}")
            return False